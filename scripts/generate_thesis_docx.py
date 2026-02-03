import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_run_fonts(run, arabic_font: str, latin_font: str, size_pt: float, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size_pt)

    # python-docx doesn't fully handle mixed RTL/LTR font mapping automatically.
    # We set a default Latin font and explicitly set the Arabic font mapping.
    run.font.name = latin_font

    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:cs"), arabic_font)
    rfonts.set(qn("w:arab"), arabic_font)


def _set_paragraph_format(paragraph, line_spacing: float = 1.15, rtl: bool = True):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing

    # Best-effort RTL. Word may still need minor adjustment depending on system.
    if rtl:
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
        bidi.set(qn("w:val"), "1")


def _add_styled_paragraph(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    arabic_font="Simplified Arabic",
    latin_font="Times New Roman",
    size_pt=13.5,
    bold=False,
    line_spacing=1.15,
    rtl=True,
):
    p = doc.add_paragraph()
    p.alignment = align
    _set_paragraph_format(p, line_spacing=line_spacing, rtl=rtl)

    # Remove common emoji / decoration at the start (keeps content readable in Word).
    cleaned = re.sub(r"^[\s\W_]+", "", text).strip()
    run = p.add_run(cleaned)
    _set_run_fonts(run, arabic_font, latin_font, size_pt=size_pt, bold=bold)
    return p


def _configure_document(doc: Document):
    section = doc.sections[0]

    # Page size: A4
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

    # Margins per guideline:
    # Right: 3 cm for binding
    # Top/Bottom/Left: 2.54 cm
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(3.0)


def _iter_markdown_lines(md_text: str):
    for raw in md_text.splitlines():
        yield raw.rstrip("\n")


def markdown_to_docx(md_path: Path, out_path: Path):
    md_text = md_path.read_text(encoding="utf-8")

    doc = Document()
    _configure_document(doc)

    in_code_block = False
    first_h1 = True

    for line in _iter_markdown_lines(md_text):
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            # Code: left align, single font, smaller size
            _add_styled_paragraph(
                doc,
                line,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                arabic_font="Times New Roman",
                latin_font="Times New Roman",
                size_pt=10.5,
                bold=False,
                line_spacing=1.0,
                rtl=False,
            )
            continue

        # Horizontal rule or separators
        if stripped in {"---", "- - -", "***"}:
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()

            if level == 1:
                # First H1 is treated as document title, later H1 as chapter title on separate page.
                if not first_h1:
                    doc.add_page_break()

                _add_styled_paragraph(
                    doc,
                    title,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    size_pt=18.0,
                    bold=True,
                    line_spacing=1.15,
                    rtl=True,
                )
                first_h1 = False
                continue

            if level == 2:
                _add_styled_paragraph(doc, title, size_pt=16.0, bold=True)
                continue

            if level == 3:
                _add_styled_paragraph(doc, title, size_pt=14.5, bold=True)
                continue

            # level >= 4
            _add_styled_paragraph(doc, title, size_pt=13.5, bold=True)
            continue

        # Bullet list
        if stripped.startswith("- "):
            text = stripped[2:].strip()
            p = _add_styled_paragraph(doc, text, size_pt=13.5, bold=False)
            # Use Word's built-in bullet formatting
            p.style = doc.styles["List Bullet"]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_format(p, line_spacing=1.15, rtl=True)
            continue

        # Numbered list (simple)
        m = re.match(r"^(\d+)\.(\s+)(.*)$", stripped)
        if m:
            text = m.group(3).strip()
            p = _add_styled_paragraph(doc, text, size_pt=13.5, bold=False)
            p.style = doc.styles["List Number"]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_format(p, line_spacing=1.15, rtl=True)
            continue

        # Normal paragraph
        _add_styled_paragraph(doc, stripped, size_pt=13.5, bold=False)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input",
        default=str(Path("docs") / "THESIS_COMPLETE.md"),
        help="Path to input markdown file",
    )
    parser.add_argument(
        "--output",
        default=str(Path("docs") / "THESIS_COMPLETE.docx"),
        help="Path to output docx file",
    )
    args = parser.parse_args()

    md_path = Path(args.input)
    out_path = Path(args.output)

    if not md_path.exists():
        raise FileNotFoundError(f"Input not found: {md_path}")

    markdown_to_docx(md_path, out_path)


if __name__ == "__main__":
    main()
